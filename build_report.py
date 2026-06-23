# -*- coding: utf-8 -*-
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSETS = "report_assets"

with open(f"{ASSETS}/results.json", "r", encoding="utf-8") as f:
    R = json.load(f)

NAVY = RGBColor(0x1F, 0x35, 0x5E)
ORANGE = RGBColor(0xD9, 0x82, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=NAVY, size=14, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "1F355E")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_para(text, size=10.5, color=None, bold=False, italic=False, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, size=10.5, space_after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_image(path, width_cm=15.5, caption=None):
    doc.add_picture(path, width=Cm(width_cm))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        cap.paragraph_format.space_after = Pt(8)


def make_table(headers, rows, widths_cm=None, header_color="1F355E"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return table


# ===================== KAPAK / BAŞLIK =====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("AMAZON KABLOSUZ KULAKLIK ÜRÜNLERİNDE\nAÇIKLANABİLİR ÜRÜN ANALİTİĞİ")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = NAVY
title_p.paragraph_format.space_after = Pt(4)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Yönetici Özeti Raporu — YBS Python ile Veri Bilimi Dönem Sonu Projesi")
run.font.size = Pt(11.5)
run.italic = True
run.font.color.rgb = ORANGE
sub_p.paragraph_format.space_after = Pt(6)

repo_p = doc.add_paragraph()
repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = repo_p.add_run("GitHub Deposu: https://github.com/elifaydnlll/amazon-kablosuz-kulaklik-urunlerinde-aciklanabilir-urun-analitigi/tree/main")
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
repo_p.paragraph_format.space_after = Pt(14)

# ===================== 1. PROBLEM TANIMI =====================
add_heading("1. Problem Tanımı ve Proje Amacı")
add_para(
    "Amazon üzerinde satılan kablosuz kulaklık ürünlerine yönelik müşteri yorumları, ürün geliştirme "
    "kararlarına yön verebilecek değerli ama yapılandırılmamış bir veri kaynağıdır. Bu projede, "
    "müşteri yorumları (metin verisi) ile ürün fiyat/indirim bilgileri (yapılandırılmış veri) "
    "birleştirilerek, müşteri memnuniyetini etkileyen temel faktörlerin veri temelli olarak ortaya "
    "konması ve ürün geliştirme yatırımlarının finansal etkisinin ölçülmesi amaçlanmıştır. Projenin "
    "nihai çıktısı, ürün yöneticilerinin “hangi ürün problemini önce çözmeliyiz?” sorusuna kanıta "
    "dayalı yanıt verebileceği bir karar destek çerçevesidir."
)

# ===================== 2. VERİ KAYNAKLARI =====================
add_heading("2. Veri Kaynakları ve Veri Harmanlama (Data Fusion)")
add_para("Proje, biri metin tabanlı diğeri yapılandırılmış olmak üzere iki farklı veri kaynağının birleştirilmesine dayanmaktadır:")
add_bullet(f"AllProductReviews.csv — {R['shape_reviews'][0]:,} satır müşteri yorumu (başlık, yorum metni, yıldız puanı, ürün adı).")
add_bullet(f"ProductInfo.csv — {R['shape_products'][0]} üründen oluşan fiyat/indirim/URL bilgisi (MRP, satış fiyatı, ürün adı).")
add_para(
    f"İki veri seti ürün adı (Product / ProductShortName) üzerinden \"left join\" ile birleştirilmiş, "
    f"sonucunda {R['merged_shape'][0]:,} satır ve {R['merged_shape'][1]} sütundan oluşan bütünleşik bir "
    "analiz tablosu elde edilmiştir. Birleştirme sonrasında eşleşmeyen kayıt bulunmadığı doğrulanmıştır; "
    "böylece her bir müşteri yorumu, ait olduğu ürünün fiyat ve indirim bilgisiyle aynı satırda "
    "analiz edilebilir hale gelmiştir."
)

# ===================== 3. ÖZELLİK MÜHENDİSLİĞİ =====================
add_heading("3. İş Mantığına Dayalı Özellik Mühendisliği")
add_para("Ham veri sütunlarının ötesinde, işletme mantığına dayanan aşağıdaki türetilmiş değişkenler oluşturulmuştur:")
add_bullet("DiscountRate — (MRP − Satış Fiyatı) / MRP formülüyle hesaplanan indirim oranı; fiyatlama stratejisinin memnuniyetle ilişkisini test etmek için.")
add_bullet("SentimentScore / SentimentLabel — VADER algoritması ile yorum metninden çıkarılan duygu skoru (-1 ile +1 arası) ve Positive/Negative/Neutral etiketi.")
add_bullet("ComplaintCategory — Yorum metnindeki anahtar kelimelere göre kural tabanlı sınıflandırma (Battery, Connectivity, Sound, Comfort, Mic, Durability, Price, Other); hangi ürün probleminin öncelikli olduğunu belirlemek için.")
add_bullet("ReviewLength — Yorum metninin karakter uzunluğu; yorum detaylılığının memnuniyetle ilişkisini ölçmek için.")
add_bullet("Satisfied (hedef değişken) — 4 ve 5 yıldız \"memnun\" (1), 1-3 yıldız \"memnun değil\" (0) olarak etiketlenmiştir.")

# ===================== 4. EDA =====================
add_heading("4. Keşifsel Veri Analizi (EDA) Bulguları")
add_para(
    f"Toplam {R['merged_shape'][0]:,} yorumun %{R['satisfied_pct']:.1f}'i memnun (4-5 yıldız) olarak "
    f"sınıflandırılmıştır. Duygu analizi etiketlerine göre yorumların %{100*R['sentiment_counts']['Positive']/R['merged_shape'][0]:.0f}'i "
    "pozitif, geri kalanı negatif/nötr olarak ayrışmıştır."
)
add_image(f"{ASSETS}/01_star_distribution.png", width_cm=8.2)
add_image(f"{ASSETS}/03_complaint_category.png", width_cm=8.2)

cat_sat = R["category_satisfaction"]
cat_rows = sorted(cat_sat.items(), key=lambda kv: -kv[1]["NotSatisfied%"])
add_para(
    "Şikayet kategorileri bazında memnuniyetsizlik oranları incelendiğinde, en kritik problem alanının "
    f"\"Mic\" (mikrofon/arama kalitesi) kategorisi olduğu görülmüştür — bu kategoriye giren yorumların "
    f"%{cat_rows[0][1]['NotSatisfied%']:.1f}'i memnuniyetsizdir. Bunu \"Connectivity\" (%{cat_rows[1][1]['NotSatisfied%']:.1f}) "
    f"ve \"Comfort\" (%{cat_rows[2][1]['NotSatisfied%']:.1f}) kategorileri takip etmektedir.",
    space_after=4
)
add_image(f"{ASSETS}/04_category_satisfaction.png", width_cm=8.2, caption="Şekil: Kategori bazlı memnuniyetsizlik oranı (%)")

# ===================== 5. KORELASYON + MODEL =====================
add_heading("5. Korelasyon Analizi ve Tahmin Modeli")
add_para(
    "Sayısal değişkenler arası korelasyon analizinde, müşteri memnuniyeti (Satisfied) ile en güçlü "
    f"ilişkinin yorum puanı (r={R['corr_matrix']['Satisfied']['ReviewStar']:.2f}, beklenen) ve duygu skoru "
    f"(r={R['corr_matrix']['Satisfied']['SentimentScore']:.2f}) arasında olduğu görülmüştür. İndirim oranı ve yorum "
    "uzunluğunun memnuniyetle ilişkisi zayıftır; bu da memnuniyetin fiyat indiriminden çok ürünün "
    "fiili performansından kaynaklandığına işaret etmektedir."
)
add_para(
    "Müşteri memnuniyetini tahmin etmek için fiyat, indirim oranı, yorum uzunluğu, duygu skoru, "
    "şikayet kategorisi ve ürün adı değişkenleriyle beslenen bir Random Forest sınıflandırma modeli "
    "kurulmuştur (200 ağaç, sınıf dengesizliği için class_weight=\"balanced\")."
)
make_table(
    ["Metrik", "Değer"],
    [
        ["Test Doğruluğu (Accuracy)", f"%{R['accuracy']*100:.1f}"],
        ["Memnun Sınıfı F1-Score", f"{R['classification_report']['1']['f1-score']:.2f}"],
        ["Memnun Değil Sınıfı F1-Score", f"{R['classification_report']['0']['f1-score']:.2f}"],
        ["Makro Ortalama F1-Score", f"{R['classification_report']['macro avg']['f1-score']:.2f}"],
    ],
    widths_cm=[8, 4]
)
add_para("")

# ===================== 6. XAI / SHAP =====================
add_heading("6. Açıklanabilir Yapay Zeka (SHAP) Bulguları")
add_para(
    "Random Forest modelinin “kara kutu” olmaktan çıkarılması için SHAP (TreeExplainer) yöntemi "
    "uygulanmıştır. SHAP analizine göre modelin memnuniyet tahminini en çok etkileyen değişken, "
    "yorum metninden türetilen duygu skoru (SentimentScore) olmuştur; bunu yorum uzunluğu, ürünün "
    "MRP/satış fiyatı ve \"Sound\" şikayet kategorisi takip etmektedir. Bu bulgu, özellik mühendisliği "
    "aşamasında metinden üretilen duygu skorunun modelin tahmin gücüne en büyük katkıyı sağladığını "
    "ve müşteri memnuniyetinin önce ürünün performansından (ses, pil, bağlantı), sonra fiyat "
    "değişkenlerinden etkilendiğini göstermektedir."
)
add_image(f"{ASSETS}/06_shap_summary.png", width_cm=14, caption="Şekil: SHAP özet grafiği — değişkenlerin “memnun” tahminine etkisi")

# ===================== 7. FİNANSAL SİMÜLASYON =====================
add_heading("7. İş Senaryosu ve Finansal Simülasyon")
add_para(
    f"EDA ve SHAP bulguları, mikrofon/arama kalitesi (\"Mic\") kategorisinin en yüksek memnuniyetsizlik "
    f"oranına (%{cat_rows[0][1]['NotSatisfied%']:.1f}) sahip olduğunu göstermiştir. Bu bulgu üzerine, mikrofon "
    "kalitesinin iyileştirilmesi durumunda satış dönüşüm oranındaki artışın yaratacağı finansal etki "
    "simüle edilmiştir."
)
make_table(
    ["Varsayım / Sonuç", "Değer"],
    [
        ["Aylık ürün sayfası ziyaretçisi", f"{100000:,}"],
        ["Mevcut dönüşüm oranı", "%3,0"],
        ["İyileştirme sonrası dönüşüm oranı", "%3,7"],
        ["Ortalama ürün fiyatı", f"{R['average_price']:,.0f} ₹"],
        ["Ek aylık satış adedi", f"{R['additional_sales']:,.0f} adet"],
        ["Ek aylık kâr", f"{R['additional_profit']:,.0f} ₹"],
        ["Geliştirme maliyeti", f"{R['development_cost']:,.0f} ₹"],
        ["Net kâr (1. ay)", f"{R['net_profit']:,.0f} ₹"],
        ["Yatırım Getirisi (ROI)", f"%{R['roi_pct']:.0f}"],
    ],
    widths_cm=[9, 6]
)
add_para(
    f"Sonuçlara göre, mikrofon problemlerinin giderilmesine yönelik 100.000 ₹'lik bir geliştirme "
    f"yatırımı, dönüşüm oranındaki %0,7 puanlık artış varsayımıyla yaklaşık {R['net_profit']:,.0f} ₹ "
    f"net kâr ve %{R['roi_pct']:.0f} ROI sağlayabilecektir. Bu, ürün geliştirme bütçesinin mikrofon "
    "iyileştirmesine ayrılmasının finansal olarak güçlü bir gerekçeye sahip olduğunu göstermektedir.",
    space_after=4
)

# ===================== 8. SONUÇ =====================
add_heading("8. Sonuç ve Öneriler")
add_bullet("Müşteri memnuniyeti en güçlü şekilde yorum duygu skoru ve şikayet kategorisiyle açıklanmaktadır; fiyat/indirim etkisi sınırlıdır.")
add_bullet("Mikrofon (Mic) ve bağlantı (Connectivity) kategorileri en kritik memnuniyetsizlik kaynaklarıdır ve öncelikli olarak ele alınmalıdır.")
add_bullet(f"Kurulan Random Forest modeli %{R['accuracy']*100:.0f} doğrulukla memnuniyeti tahmin edebilmekte ve SHAP ile açıklanabilir hale getirilmiştir.")
add_bullet(f"Mikrofon iyileştirmesine yapılacak yatırımın yaklaşık %{R['roi_pct']:.0f} ROI ile geri dönmesi beklenmektedir; bu nedenle ürün yol haritasında önceliklendirilmesi önerilir.")
add_bullet("Gelecek çalışmalarda, ek metin madenciliği teknikleri (konu modelleme, BERT tabanlı duygu analizi) ile ComplaintCategory sınıflandırmasının doğruluğu artırılabilir.")

add_para("")
add_para("Proje GitHub Deposu:", bold=True, size=11, color=NAVY)
add_para("https://github.com/elifaydnlll/amazon-kablosuz-kulaklik-urunlerinde-aciklanabilir-urun-analitigi/tree/main", size=10, italic=True)

doc.save("Yonetici_Ozeti_Raporu.docx")
print("Rapor oluşturuldu: Yonetici_Ozeti_Raporu.docx")
